"""
parser.py
Extracts raw text from uploaded resume files (PDF / DOCX) and provides
simple heuristics to detect which professional domain a resume belongs to.
"""

import os
import pdfplumber
import docx

from utils import SKILL_BANK, clean_text


def extract_text_from_pdf(filepath: str) -> str:
    text_parts = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    return clean_text("\n".join(text_parts))


def extract_text_from_docx(filepath: str) -> str:
    document = docx.Document(filepath)
    paragraphs = [p.text for p in document.paragraphs]

    # Also grab text sitting inside tables (common in resume templates)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)

    return clean_text("\n".join(paragraphs))


def extract_text(filepath: str) -> str:
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(filepath)
    elif ext == ".docx":
        return extract_text_from_docx(filepath)
    raise ValueError(f"Unsupported file type: {ext}")


def detect_domain(resume_text: str) -> str:
    """Guess the resume's professional domain by counting skill-bank hits."""
    lower = resume_text.lower()
    best_domain = "general"
    best_score = 0

    for domain, skills in SKILL_BANK.items():
        if domain == "general":
            continue
        score = sum(1 for skill in skills if skill in lower)
        if score > best_score:
            best_score = score
            best_domain = domain

    return best_domain if best_score > 0 else "general"
